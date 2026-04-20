#!/usr/bin/env python3
"""
金蝶文档写作助手 - Markdown 转 Word 导出脚本

使用方式：
    python scripts/export-word.py input.md output.docx

功能：
    - 自动检查并安装 python-docx 依赖
    - 将 Markdown 内容转换为 Word 文档
    - 支持标题层级、加粗、斜体、列表、表格、代码块、引用块
    - 生成后验证文件是否存在并输出文件大小
"""

import sys
import os
import subprocess
import re


def ensure_docx_installed():
    """检查并安装 python-docx 依赖"""
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        print("[INFO] python-docx 未安装，正在自动安装...")
        try:
            subprocess.check_call(
                [sys.executable, "-m", "pip", "install", "python-docx"],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
            )
            print("[INFO] python-docx 安装成功")
            return True
        except subprocess.CalledProcessError as e:
            print(f"[ERROR] 安装 python-docx 失败: {e}")
            return False


def parse_inline(paragraph, text):
    """解析行内格式（加粗、斜体）并添加到段落"""
    # 匹配 ***bold italic***、**bold**、*italic*、`code`
    patterns = [
        (r"\*\*\*(.+?)\*\*\*", {"bold": True, "italic": True}),
        (r"\*\*(.+?)\*\*", {"bold": True}),
        (r"\*(.+?)\*", {"italic": True}),
        (r"`(.+?)`", {"font_name": "Courier New"}),
    ]

    # 合并所有模式用于分割
    combined = r"(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*|`.+?`)"
    parts = re.split(combined, text)

    for part in parts:
        if not part:
            continue
        matched = False
        for pattern, style in patterns:
            m = re.fullmatch(pattern, part)
            if m:
                run = paragraph.add_run(m.group(1))
                if style.get("bold"):
                    run.bold = True
                if style.get("italic"):
                    run.italic = True
                if style.get("font_name"):
                    run.font.name = "Courier New"
                matched = True
                break
        if not matched:
            paragraph.add_run(part)


def convert_markdown_to_docx(md_path, docx_path):
    """将 Markdown 文件转换为 Word 文档"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if not os.path.exists(md_path):
        print(f"[ERROR] 输入文件不存在: {md_path}")
        return False

    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    lines = content.split("\n")
    doc = Document()

    i = 0
    while i < len(lines):
        line = lines[i]

        # 代码块
        if line.strip().startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # 跳过结束的 ```
            # 添加代码块为带灰色背景的段落
            for code_line in code_lines:
                p = doc.add_paragraph()
                run = p.add_run(code_line)
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Pt(18)
            continue

        # 引用块
        if line.strip().startswith(">"):
            quote_lines = []
            while i < len(lines) and (lines[i].strip().startswith(">") or lines[i].strip() == ""):
                if lines[i].strip() == "" and (i + 1 >= len(lines) or not lines[i + 1].strip().startswith(">")):
                    break
                quote_text = re.sub(r"^>\s?", "", lines[i])
                quote_lines.append(quote_text)
                i += 1
            quote_content = "\n".join(quote_lines)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(36)
            run = p.add_run(quote_content)
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            continue

        # 表格
        if "|" in line and i + 1 < len(lines) and re.match(r"^\s*\|[\s\-:|]+\|\s*$", lines[i + 1]):
            table_lines = []
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i])
                i += 1
            if len(table_lines) >= 2:
                # 解析表头
                headers = [cell.strip() for cell in table_lines[0].split("|")[1:-1]]
                # 跳过分隔行，解析数据行
                rows_data = []
                for tl in table_lines[2:]:
                    cells = [cell.strip() for cell in tl.split("|")[1:-1]]
                    rows_data.append(cells)

                num_cols = len(headers)
                table = doc.add_table(rows=1 + len(rows_data), cols=num_cols)
                table.style = "Table Grid"

                # 填充表头
                for j, header in enumerate(headers):
                    if j < num_cols:
                        table.rows[0].cells[j].text = header

                # 填充数据
                for row_idx, row_data in enumerate(rows_data):
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < num_cols:
                            table.rows[row_idx + 1].cells[col_idx].text = cell_text
            continue

        # 标题
        heading_match = re.match(r"^(#{1,5})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            doc.add_heading(text, level=level)
            i += 1
            continue

        # 无序列表
        ul_match = re.match(r"^(\s*)([-*+])\s+(.+)$", line)
        if ul_match:
            indent = len(ul_match.group(1))
            text = ul_match.group(3)
            p = doc.add_paragraph(style="List Bullet")
            if indent >= 4:
                p.paragraph_format.left_indent = Pt(36)
            elif indent >= 2:
                p.paragraph_format.left_indent = Pt(18)
            parse_inline(p, text)
            i += 1
            continue

        # 有序列表
        ol_match = re.match(r"^(\s*)\d+[.)]\s+(.+)$", line)
        if ol_match:
            indent = len(ol_match.group(1))
            text = ol_match.group(2)
            p = doc.add_paragraph(style="List Number")
            if indent >= 4:
                p.paragraph_format.left_indent = Pt(36)
            elif indent >= 2:
                p.paragraph_format.left_indent = Pt(18)
            parse_inline(p, text)
            i += 1
            continue

        # 空行
        if line.strip() == "":
            i += 1
            continue

        # 水平分隔线
        if re.match(r"^---+$|^\*\*\*+$|^___+$", line.strip()):
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        parse_inline(p, line)
        i += 1

    # 确保输出目录存在
    output_dir = os.path.dirname(docx_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)

    # 保存文档
    doc.save(docx_path)
    return True


def main():
    if len(sys.argv) != 3:
        print("使用方式: python scripts/export-word.py <input.md> <output.docx>")
        print("示例: python scripts/export-word.py article.md output/article.docx")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]

    print(f"[INFO] 输入文件: {input_path}")
    print(f"[INFO] 输出文件: {output_path}")

    # 检查并安装依赖
    if not ensure_docx_installed():
        print("[ERROR] 无法安装 python-docx，请手动运行: pip install python-docx")
        sys.exit(1)

    # 转换
    print("[INFO] 正在转换 Markdown 到 Word...")
    success = convert_markdown_to_docx(input_path, output_path)

    if not success:
        print("[ERROR] 转换失败")
        sys.exit(1)

    # 验证文件
    if os.path.exists(output_path):
        file_size = os.path.getsize(output_path)
        if file_size > 0:
            print(f"[SUCCESS] Word 文档生成成功!")
            print(f"[SUCCESS] 文件路径: {os.path.abspath(output_path)}")
            print(f"[SUCCESS] 文件大小: {file_size:,} 字节")
        else:
            print("[ERROR] 文件已生成但大小为 0，可能转换过程出错")
            sys.exit(1)
    else:
        print(f"[ERROR] 文件未生成: {output_path}")
        sys.exit(1)


if __name__ == "__main__":
    main()
